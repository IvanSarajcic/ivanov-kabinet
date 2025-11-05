"""
Скрипт за екстракцију чистог текста из .docx фајла
"""

from docx import Document
import sys

def extract_text_from_docx(docx_path, output_path=None):
    """
    Извлачи сав текст из .docx фајла без форматирања.
    
    Args:
        docx_path: путања до .docx фајла
        output_path: путања до излазног .txt фајла (опционо)
    """
    try:
        # Учитај документ
        doc = Document(docx_path)
        
        # Извуци текст из свих параграфа
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Прескочи празне линије
                full_text.append(paragraph.text)
        
        # Извуци текст из табела
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        # Споји све у један текст
        result = '\n'.join(full_text)
        
        # Сачувај у фајл ако је наведена путања
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result)
            print(f"Текст сачуван у: {output_path}")
        
        return result
    
    except Exception as e:
        print(f"Грешка: {e}")
        return None

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Употреба: python extract_text.py <путања_до_docx> [излазни_фајл.txt]")
        print("Пример: python extract_text.py plan.docx plan.txt")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else docx_file.replace('.docx', '.txt')
    
    text = extract_text_from_docx(docx_file, output_file)
    
    if text:
        print(f"\nИзвучено {len(text)} карактера")
        print("\nПрвих 500 карактера:")
        print("-" * 50)
        print(text[:500])
