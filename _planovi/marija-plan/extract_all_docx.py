"""
Скрипт за екстракцију текста из свих .docx фајлова у фолдеру
"""

import os
from docx import Document

def extract_all_docx_in_folder(folder_path='.'):
    """Извлачи текст из свих .docx фајлова у датом фолдеру"""
    
    # Пронађи све .docx фајлове
    docx_files = [f for f in os.listdir(folder_path) 
                  if f.endswith('.docx') and not f.startswith('~$')]
    
    if not docx_files:
        print("Нема .docx фајлова у фолдеру")
        return
    
    print(f"Пронађено {len(docx_files)} .docx фајлова\n")
    
    for docx_file in sorted(docx_files):
        print(f"Обрађујем: {docx_file}")
        
        try:
            # Учитај документ
            doc = Document(os.path.join(folder_path, docx_file))
            
            # Извуци текст
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Извуци из табела
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            result = '\n'.join(full_text)
            
            # Креирај име излазног фајла
            output_file = docx_file.replace('.docx', '.txt')
            output_path = os.path.join(folder_path, output_file)
            
            # Сачувај
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result)
            
            print(f"  ✓ Сачувано у: {output_file} ({len(result)} карактера)\n")
        
        except Exception as e:
            print(f"  ✗ Грешка: {e}\n")

if __name__ == "__main__":
    import sys
    
    folder = sys.argv[1] if len(sys.argv) > 1 else '.'
    
    print(f"Екстракција текста из фолдера: {os.path.abspath(folder)}\n")
    print("=" * 60)
    
    extract_all_docx_in_folder(folder)
    
    print("=" * 60)
    print("Готово!")
