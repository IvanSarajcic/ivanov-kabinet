"""
Скрипт за екстракцију чистог текста из .doc и .docx фајла
"""

import sys
import os

def extract_from_docx(docx_path):
    """Извлачи текст из .docx фајла"""
    from docx import Document
    doc = Document(docx_path)
    full_text = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    
    return '\n'.join(full_text)

def extract_from_doc(doc_path):
    """Извлачи текст из старог .doc формата"""
    try:
        import win32com.client as win32
        
        # Креирај Word апликацију
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        
        # Отвори документ
        doc = word.Documents.Open(os.path.abspath(doc_path))
        
        # Извуци текст
        text = doc.Content.Text
        
        # Затвори документ и апликацију
        doc.Close(False)
        word.Quit()
        
        return text
    
    except ImportError:
        print("Инсталирај pywin32: pip install pywin32")
        return None
    except Exception as e:
        print(f"Грешка при читању .doc фајла: {e}")
        return None

def extract_text(file_path, output_path=None):
    """Извлачи текст из .doc или .docx фајла"""
    
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == '.docx':
            text = extract_from_docx(file_path)
        elif ext == '.doc':
            text = extract_from_doc(file_path)
        else:
            print(f"Неподржан формат: {ext}")
            return None
        
        if text and output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"Текст сачуван у: {output_path}")
        
        return text
    
    except Exception as e:
        print(f"Грешка: {e}")
        return None

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Употреба: python extract_text_doc.py <путања_до_фајла> [излазни_фајл.txt]")
        print("Пример: python extract_text_doc.py plan.doc plan.txt")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.rsplit('.', 1)[0] + '.txt'
    
    text = extract_text(input_file, output_file)
    
    if text:
        print(f"\nИзвучено {len(text)} карактера")
        print("\nПрвих 500 карактера:")
        print("-" * 50)
        print(text[:500])
